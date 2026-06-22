from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_all_font_styles(doc, font_name='宋体', font_size=12, color_rgb=(0, 0, 0)):
    """
    强制修改文档中所有文字的字体、大小和颜色
    :param doc_path: 原文件路径
    :param save_path: 保存路径
    :param font_name: 中文字体名称 (如 '宋体', '黑体')
    :param font_size: 字号 (磅值 pt)
    :param color_rgb: 颜色元组 (R, G, B)
    """
    target_color = RGBColor(*color_rgb)

    # --- 1. 修改正文段落 ---
    for paragraph in doc.paragraphs:
        # 关键步骤：设置段落格式为两端对齐
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.color.rgb = target_color
            # 关键：设置东亚字体（解决中文不生效的问题）
            rPr = run._r.get_or_add_rPr()
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:eastAsia'), font_name)
            rPr.append(rFonts)

    # --- 2. 修改表格内的文字 ---
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = font_name
                        run.font.size = Pt(font_size)
                        run.font.color.rgb = target_color
                        # 同样处理东亚字体
                        rPr = run._r.get_or_add_rPr()
                        rFonts = OxmlElement('w:rFonts')
                        rFonts.set(qn('w:eastAsia'), font_name)
                        rPr.append(rFonts)

                    # 关键步骤：设置段落格式为两端对齐
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


                    # 获取段落的 pPr 元素
                    pPr = paragraph._p.get_or_add_pPr()
                    # 创建 spacing 元素并设置属性
                    spacing = OxmlElement('w:spacing')
                    spacing.set(qn('w:lineRule'), 'auto')       # 自动行距规则
                    spacing.set(qn('w:line'), '360')            # 360 = 1.5 * 240 (Word 单位)

                    # 将 spacing 添加到 pPr
                    pPr.append(spacing)

    # --- 3. 修改页眉和页脚 ---
    for section in doc.sections:
        for header_footer in [section.header, section.footer]:
            for paragraph in header_footer.paragraphs:
                for run in paragraph.runs:
                    run.font.name = font_name
                    run.font.size = Pt(font_size)
                    run.font.color.rgb = target_color
                    rPr = run._r.get_or_add_rPr()
                    rFonts = OxmlElement('w:rFonts')
                    rFonts.set(qn('w:eastAsia'), font_name)
                    rPr.append(rFonts)

    print(f"处理完成")
