import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn

from tools.get_heading_level import get_heading_level, heading_styles
from tools.set_table import set_table_width_to_100_percent, set_table_to_embedded, set_cell_margins, set_cell_format, set_table_borders

# ==========================================
# 1. 辅助与工具函数
# ==========================================

def get_style_name(para):
    """
    安全获取段落样式名称
    防止 para.style 为 None 导致报错
    """
    if para.style is not None:
        return para.style.name
    return ""

def generate_output_path(input_file, output_file=None):
    """生成输出文件路径"""
    if output_file:
        return output_file

    directory = os.path.dirname(input_file)
    filename = os.path.basename(input_file)
    name_only, _ = os.path.splitext(filename)

    # 默认输出为 docx，如果需要转 pdf 可以在这里改后缀
    new_filename = f"{name_only}-formated.docx"
    return os.path.join(directory, new_filename)

# ==========================================
# 2. 格式化核心逻辑
# ==========================================

def format_all_paragraphs(doc):
    """
    遍历所有段落，根据样式应用格式
    """
    for para in doc.paragraphs:
        style_name = get_style_name(para)

        # --- 标题处理 ---
        heading_level = get_heading_level(style_name)
        if heading_level:
            # 2. 为不同级别的标题定义不同的样式配置
            # 你可以在这里随意修改每一级的样式


            # 3. 获取当前级别的样式配置，如果未定义则使用默认配置
            style_config = heading_styles.get(heading_level, {
                'font_size': Pt(12),
                'font_color': RGBColor(0, 0, 0),
                'bg_color': None,
                'alignment': WD_PARAGRAPH_ALIGNMENT.LEFT,
                'first_line_indent': None
            })

            # 4. 应用段落格式
            para.paragraph_format.alignment = style_config['alignment']
            para.paragraph_format.first_line_indent = style_config['first_line_indent']
            
            # 注意：python-docx 设置段落背景色比较复杂，通常不直接支持。
            # 如果需要背景色，可能需要更底层的 XML 操作，这里暂不展开。

            # 5. 应用字体格式
            for run in para.runs:
                run.font.name = 'SimHei' # 统一用黑体，也可以根据级别再区分
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
                run.font.bold = True
                run.font.size = style_config['font_size']
                run.font.color.rgb = style_config['font_color']

        # --- 正文处理 ---
        else:

            # 1. 设置段落格式
            para_format = para.paragraph_format

            # 设置 1.5 倍行距
            para_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

            # 设置首行缩进 2 字符
            # 正文字号通常为 12pt (小四)，2字符即 24pt
            para_format.first_line_indent = Pt(24)

            # 设置段前段后间距（可选，例如 0 行）
            para_format.space_before = Pt(0)
            para_format.space_after = Pt(0)


            # 下面的代码不管用
            pPr = para._p.get_or_add_pPr()
            # 强制取消 "与网格对齐" (snapToGrid)
            pPr.set(qn('w:snapToGrid'), '0')
            # 强制取消 "自动调整右缩进" (autoAdjustRightIndent)
            pPr.set(qn('w:autoAdjustRightIndent'), '0')

            # 2. 设置字体格式
            for run in para.runs:
                # 设置中文字体为宋体，英文字体为 Times New Roman
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)


    return doc

def format_page_layout(doc):
    """
    设置页面边距
    """
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    return doc

# ==========================================
# 3. 主流程
# ==========================================

def process_document(input_file, output_file=None):
    """
    主处理函数
    """
    print(f"正在处理文档：{input_file} ...")

    # 1. 加载文档
    doc = Document(input_file)

    # 2. 设置页面布局
    # doc = format_page_layout(doc)

    # 3. 格式化所有段落（包含首行缩进逻辑）
    doc = format_all_paragraphs(doc)

    # 设置表格样式
    for table in doc.tables:
        set_table_width_to_100_percent(table)
        set_table_to_embedded(table)
        set_table_borders(table)

        for row in table.rows:
            for cell in row.cells:
                set_cell_margins(cell)  # 设置内边距
                set_cell_format(cell)

    # 4. 生成输出路径
    output_path = generate_output_path(input_file, output_file)

    # 5. 保存文档
    doc.save(output_path)
    print(f"文档处理完成，已保存至：{output_path}")

# ==========================================
# 4. 入口点
# ==========================================

if __name__ == "__main__":
    # 这里替换成你要处理的文件路径
    input_path = "/Users/teacher/Desktop/手册排版/单品销售手册.docx"

    # 如果第二个参数不填，会自动生成 【原文件名_1】.docx
    process_document(input_path)



